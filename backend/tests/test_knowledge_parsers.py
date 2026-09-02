from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from openpyxl import Workbook

from orion.contracts import RuntimeScope
from orion.knowledge.parsers import CompositeDocumentParser


def _pdf_bytes(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode("ascii"))
        payload.extend(obj)
        payload.extend(b"\nendobj\n")
    xref = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(payload)


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Capacity", level=1)
    document.add_paragraph("Peak requirement is forty workers.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "HCMC"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Sizing"
    worksheet["A1"] = "workers"
    worksheet["B1"] = 40
    worksheet["A2"] = "region"
    worksheet["B2"] = "HCMC"
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _scope(session_id: str, attachment_id: str) -> RuntimeScope:
    return RuntimeScope(
        session_id=session_id,
        attachment_ids=(attachment_id,),
        principal_id="local",
        workspace_id="local",
    )


def test_composite_parser_supports_pdf_with_page_provenance() -> None:
    parsed = CompositeDocumentParser().parse(_pdf_bytes("PDF sentinel 7711"), "application/pdf")

    assert "PDF sentinel 7711" in parsed.text
    assert parsed.sections[0].page == 1
    assert parsed.sections[0].section is None


def test_composite_parser_supports_docx_heading_paragraph_and_table_provenance() -> None:
    parsed = CompositeDocumentParser().parse(
        _docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "Peak requirement is forty workers." in parsed.text
    assert any(section.section == "Capacity" for section in parsed.sections)
    assert any(
        section.section and "Capacity · paragraph" in section.section
        for section in parsed.sections
    )
    assert any(section.section == "Table 1 · row 1" for section in parsed.sections)


def test_composite_parser_supports_xlsx_sheet_and_row_provenance() -> None:
    parsed = CompositeDocumentParser().parse(
        _xlsx_bytes(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    assert "A1=workers" in parsed.text
    assert "B2=HCMC" in parsed.text
    assert {section.section for section in parsed.sections} == {
        "Sheet Sizing · row 1",
        "Sheet Sizing · row 2",
    }


def test_composite_parser_can_sniff_supported_binary_formats() -> None:
    parser = CompositeDocumentParser()

    assert parser.parse(_pdf_bytes("sniff pdf"), "application/octet-stream").sections[0].page == 1
    assert "forty workers" in parser.parse(_docx_bytes(), None).text
    assert "workers" in parser.parse(_xlsx_bytes(), "application/octet-stream").text


def test_malformed_pdf_and_unsafe_office_archive_fail_closed() -> None:
    parser = CompositeDocumentParser()
    with pytest.raises(ValueError, match="PDF"):
        parser.parse(b"%PDF-not-valid", "application/pdf")

    archive = BytesIO()
    with ZipFile(archive, "w", ZIP_DEFLATED) as package:
        package.writestr("[Content_Types].xml", b"types")
        package.writestr("word/document.xml", b"document")
        package.writestr("../escape.xml", b"escape")
    with pytest.raises(ValueError, match="unsafe archive path"):
        parser.parse(archive.getvalue(), "application/octet-stream")


def test_pdf_provenance_reaches_knowledge_source_ref(knowledge, store) -> None:  # type: ignore[no-untyped-def]
    session = store.create_session()
    upload = knowledge.attach(session, "facts.pdf", _pdf_bytes("Citation page sentinel"), "application/pdf")
    assert upload.status == "ready"

    read = knowledge.read(_scope(session, upload.attachment_id), upload.document.document_id)
    source = knowledge.source_for_segment(read.segments[0])

    assert read.segments[0].page == 1
    assert source.page == 1
    assert source.document_id == upload.document.document_id


def test_docx_and_xlsx_project_documents_remain_project_scoped(knowledge, store) -> None:  # type: ignore[no-untyped-def]
    project_a = store.create_project("A")
    project_b = store.create_project("B")
    docx = knowledge.attach_project(
        str(project_a["project_id"]),
        "requirements.docx",
        _docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    xlsx = knowledge.attach_project(
        str(project_b["project_id"]),
        "sizing.xlsx",
        _xlsx_bytes(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    session_a = store.create_session(project_id=str(project_a["project_id"]))
    scope_a = RuntimeScope(
        session_id=session_a,
        project_id=str(project_a["project_id"]),
        principal_id="local",
        workspace_id="local",
    )

    assert docx.status == "ready" and xlsx.status == "ready"
    assert knowledge.search(scope_a, "forty workers", 5)[0].document.document_id == docx.document.document_id
    with pytest.raises(PermissionError):
        knowledge.search(scope_a, "workers", 5, (xlsx.document.document_id,))
