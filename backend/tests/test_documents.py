from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font

from orion.integrations.documents import DocumentCodecError, edit_document, read_document


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Unchanged paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "First cell"
    table.cell(0, 1).text = "Unchanged cell"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Plan"
    worksheet["A1"] = "Original"
    worksheet["B1"] = 2
    worksheet["B1"].font = Font(bold=True)
    worksheet["C1"] = "=B1*2"
    workbook.create_sheet("Archive")["A1"] = "Saved"
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def test_docx_read_and_structured_edits_preserve_unrelated_content() -> None:
    payload = _docx_bytes()
    read = read_document("/tmp/report.docx", payload, limit=10)
    edited, _ = edit_document(
        "/tmp/report.docx",
        payload,
        [
            {"kind": "set_paragraph", "paragraph_index": 0, "text": "Updated paragraph"},
            {
                "kind": "set_table_cell",
                "table_index": 0,
                "row": 0,
                "column": 0,
                "text": "Updated cell",
            },
        ],
    )
    document = Document(BytesIO(edited))

    assert read["format"] == "docx"
    assert {item["kind"] for item in read["items"]} == {"paragraph", "table_cell"}
    assert document.paragraphs[0].text == "Updated paragraph"
    assert document.paragraphs[1].text == "Unchanged paragraph"
    assert document.tables[0].cell(0, 0).text == "Updated cell"
    assert document.tables[0].cell(0, 1).text == "Unchanged cell"


def test_docx_corruption_is_controlled() -> None:
    with pytest.raises(DocumentCodecError, match="invalid or corrupt"):
        read_document("/tmp/report.docx", b"not a zip")


def test_xlsx_sheet_listing_range_read_and_edits_preserve_formula_and_style() -> None:
    payload = _xlsx_bytes()
    listing = read_document("/tmp/plan.xlsx", payload)
    cells = read_document("/tmp/plan.xlsx", payload, sheet="Plan", cell_range="A1:B1")
    edited, _ = edit_document(
        "/tmp/plan.xlsx",
        payload,
        [
            {"kind": "set_cell", "sheet": "Plan", "cell": "A1", "value": "Updated"},
            {"kind": "set_formula", "sheet": "Plan", "cell": "C1", "formula": "=B1*2"},
        ],
    )
    workbook = load_workbook(BytesIO(edited), data_only=False)

    assert [item["name"] for item in listing["sheets"]] == ["Plan", "Archive"]
    assert cells["cells"] == [{"cell": "A1", "value": "Original"}, {"cell": "B1", "value": 2}]
    assert workbook["Plan"]["A1"].value == "Updated"
    assert workbook["Plan"]["B1"].font.bold is True
    assert workbook["Plan"]["C1"].value == "=B1*2"
    assert workbook["Archive"]["A1"].value == "Saved"


@pytest.mark.parametrize("path", ["/tmp/legacy.doc", "/tmp/legacy.xls", "/tmp/macro.docm"])
def test_legacy_document_formats_are_rejected(path: str) -> None:
    with pytest.raises(DocumentCodecError) as error:
        read_document(path, b"anything")
    assert error.value.code == "unsupported_format"


def test_text_edits_require_exact_single_match_and_noop_is_detectable() -> None:
    with pytest.raises(DocumentCodecError) as missing:
        edit_document(
            "/tmp/file.txt",
            b"text",
            [{"kind": "replace_text", "old_text": "no", "new_text": "yes"}],
        )
    with pytest.raises(DocumentCodecError) as ambiguous:
        edit_document(
            "/tmp/file.txt",
            b"same same",
            [{"kind": "replace_text", "old_text": "same", "new_text": "new"}],
        )
    unchanged, _ = edit_document(
        "/tmp/file.txt", b"same", [{"kind": "replace_text", "old_text": "same", "new_text": "same"}]
    )

    assert missing.value.code == "not_found"
    assert ambiguous.value.code == "ambiguous"
    assert unchanged == b"same"


def test_office_semantic_noops_return_original_bytes_without_destroying_formatting() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Bold").bold = True
    paragraph.add_run(" text")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Same cell"
    output = BytesIO()
    document.save(output)
    docx_payload = output.getvalue()
    same_paragraph, _ = edit_document(
        "/tmp/report.docx",
        docx_payload,
        [{"kind": "set_paragraph", "paragraph_index": 0, "text": "Bold text"}],
    )
    same_cell, _ = edit_document(
        "/tmp/report.docx",
        docx_payload,
        [{"kind": "set_table_cell", "table_index": 0, "row": 0, "column": 0, "text": "Same cell"}],
    )
    xlsx_payload = _xlsx_bytes()
    same_value, _ = edit_document(
        "/tmp/plan.xlsx",
        xlsx_payload,
        [{"kind": "set_cell", "sheet": "Plan", "cell": "A1", "value": "Original"}],
    )
    same_formula, _ = edit_document(
        "/tmp/plan.xlsx",
        xlsx_payload,
        [{"kind": "set_formula", "sheet": "Plan", "cell": "C1", "formula": "=B1*2"}],
    )

    assert same_paragraph == docx_payload and same_cell == docx_payload
    assert Document(BytesIO(same_paragraph)).paragraphs[0].runs[0].bold is True
    assert same_value == xlsx_payload and same_formula == xlsx_payload


def test_office_archives_and_out_of_bounds_coordinates_are_controlled() -> None:
    archive = BytesIO()
    with ZipFile(archive, "w", ZIP_DEFLATED) as package:
        package.writestr("word/document.xml", b"x" * (8 * 1024 * 1024 + 1))
    with pytest.raises(DocumentCodecError) as expanded:
        read_document("/tmp/large.docx", archive.getvalue())
    macro = BytesIO()
    with ZipFile(macro, "w") as package:
        package.writestr("word/vbaProject.bin", b"macro")
    with pytest.raises(DocumentCodecError) as macro_error:
        read_document("/tmp/renamed.docx", macro.getvalue())
    member_count = BytesIO()
    with ZipFile(member_count, "w") as package:
        for index in range(1001):
            package.writestr(f"word/entry-{index}.xml", b"x")
    with pytest.raises(DocumentCodecError) as member_error:
        read_document("/tmp/members.docx", member_count.getvalue())
    payload = _xlsx_bytes()
    for cell_range in ("XFE1", "A1048577", "ZZZ9999999"):
        with pytest.raises(DocumentCodecError) as bounds:
            read_document("/tmp/plan.xlsx", payload, sheet="Plan", cell_range=cell_range)
        assert bounds.value.code == "invalid_input"
    with pytest.raises(DocumentCodecError) as edit_bounds:
        edit_document(
            "/tmp/plan.xlsx",
            payload,
            [{"kind": "set_cell", "sheet": "Plan", "cell": "XFE1", "value": "x"}],
        )
    boundary, boundary_summary = edit_document(
        "/tmp/plan.xlsx",
        payload,
        [{"kind": "set_cell", "sheet": "Plan", "cell": "XFD1048576", "value": "edge"}],
    )

    assert expanded.value.code == "too_large"
    assert macro_error.value.code == "unsupported_content"
    assert member_error.value.code == "too_large"
    assert edit_bounds.value.code == "invalid_input"
    assert boundary != payload and boundary_summary["changed"] is True
