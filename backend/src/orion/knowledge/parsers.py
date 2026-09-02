"""Safe local parsers for Project and session knowledge documents."""

from __future__ import annotations

from io import BytesIO
from pathlib import PurePosixPath
from zipfile import BadZipFile, ZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from orion.integrations.documents import DocumentCodecError, _office_zip_preflight
from orion.knowledge.local import PlainTextParser
from orion.knowledge.ports import ParsedDocument, ParsedSection
from orion.paths import document_upload_limit

_TEXT_MEDIA_TYPES = {"text/plain", "text/markdown", "text/x-markdown"}
_PDF_MEDIA_TYPES = {"application/pdf"}
_DOCX_MEDIA_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_XLSX_MEDIA_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
_GENERIC_MEDIA_TYPES = {None, "", "application/octet-stream"}
_MAX_EXTRACTED_CHARACTERS = 2_000_000
_MAX_PDF_PAGES = 500
_MAX_XLSX_CELLS = 100_000
_MAX_ZIP_COMPRESSION_RATIO = 1_000


class CompositeDocumentParser:
    """Parse supported local knowledge formats without changing retrieval contracts."""

    def __init__(self, maximum_bytes: int | None = None) -> None:
        self._maximum_bytes = maximum_bytes or document_upload_limit()
        self._text = PlainTextParser()

    def parse(self, content: bytes, media_type: str | None) -> ParsedDocument:
        if len(content) > self._maximum_bytes:
            raise ValueError("Document exceeds Orion's configured upload limit")
        normalized = media_type.split(";", 1)[0].strip().lower() if media_type else media_type
        kind = self._document_kind(content, normalized)
        if kind == "text":
            text_type = normalized if normalized in _TEXT_MEDIA_TYPES else None
            return self._text.parse(content, text_type)
        if kind == "pdf":
            return self._parse_pdf(content)
        if kind == "docx":
            return self._parse_docx(content)
        if kind == "xlsx":
            return self._parse_xlsx(content)
        raise ValueError(f"Unsupported document media type: {media_type}")

    def _document_kind(self, content: bytes, media_type: str | None) -> str:
        if media_type in _TEXT_MEDIA_TYPES:
            return "text"
        if media_type in _PDF_MEDIA_TYPES:
            return "pdf"
        if media_type in _DOCX_MEDIA_TYPES:
            return "docx"
        if media_type in _XLSX_MEDIA_TYPES:
            return "xlsx"
        if media_type not in _GENERIC_MEDIA_TYPES:
            raise ValueError(f"Unsupported document media type: {media_type}")
        if content.startswith(b"%PDF-"):
            return "pdf"
        if content.startswith(b"PK"):
            return self._office_kind(content)
        try:
            content.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise ValueError("Unsupported binary document content") from error
        return "text"

    @staticmethod
    def _office_kind(content: bytes) -> str:
        _validate_office_archive(content)
        try:
            with ZipFile(BytesIO(content)) as archive:
                names = {name.lower() for name in archive.namelist()}
        except (BadZipFile, OSError, ValueError) as error:
            raise ValueError("Office document is invalid or corrupt") from error
        if any(name.startswith("word/") for name in names):
            return "docx"
        if any(name.startswith("xl/") for name in names):
            return "xlsx"
        raise ValueError("Unsupported Office document package")

    @staticmethod
    def _parse_pdf(content: bytes) -> ParsedDocument:
        try:
            reader = PdfReader(BytesIO(content), strict=False)
            if reader.is_encrypted:
                raise ValueError("Encrypted PDF documents are unsupported")
            if len(reader.pages) > _MAX_PDF_PAGES:
                raise ValueError("PDF document has too many pages")
            sections: list[ParsedSection] = []
            total_characters = 0
            for page_number, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if not text:
                    continue
                total_characters += len(text)
                if total_characters > _MAX_EXTRACTED_CHARACTERS:
                    raise ValueError("PDF extracted text exceeds Orion's safety limit")
                sections.append(ParsedSection(text=text, page=page_number))
        except ValueError:
            raise
        except (PdfReadError, OSError, EOFError) as error:
            raise ValueError("Unsupported malformed PDF: invalid or corrupt") from error
        except Exception as error:  # pypdf normalizes malformed stream/XML-like internals variably.
            raise ValueError("Unsupported malformed PDF: invalid or corrupt") from error
        return _parsed_document(sections)

    @staticmethod
    def _parse_docx(content: bytes) -> ParsedDocument:
        _validate_office_archive(content)
        try:
            document = Document(BytesIO(content))
        except Exception as error:  # python-docx exposes several zip/XML exception types.
            raise ValueError("Word document is invalid or corrupt") from error
        sections: list[ParsedSection] = []
        heading: str | None = None
        total_characters = 0
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.lower().startswith("heading"):
                heading = text
                label = heading
            else:
                label = f"{heading} · paragraph {index}" if heading else f"Paragraph {index}"
            total_characters += len(text)
            if total_characters > _MAX_EXTRACTED_CHARACTERS:
                raise ValueError("Word extracted text exceeds Orion's safety limit")
            sections.append(ParsedSection(text=text, section=label))
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                values = [cell.text.strip() for cell in row.cells]
                text = " | ".join(value for value in values if value)
                if not text:
                    continue
                total_characters += len(text)
                if total_characters > _MAX_EXTRACTED_CHARACTERS:
                    raise ValueError("Word extracted text exceeds Orion's safety limit")
                sections.append(
                    ParsedSection(
                        text=text,
                        section=f"Table {table_index} · row {row_index}",
                    )
                )
        return _parsed_document(sections)

    @staticmethod
    def _parse_xlsx(content: bytes) -> ParsedDocument:
        _validate_office_archive(content)
        try:
            workbook = load_workbook(
                BytesIO(content),
                read_only=True,
                data_only=False,
                keep_vba=False,
            )
        except Exception as error:  # openpyxl exposes several zip/XML exception types.
            raise ValueError("Excel workbook is invalid or corrupt") from error
        sections: list[ParsedSection] = []
        cells_seen = 0
        total_characters = 0
        try:
            for worksheet in workbook.worksheets:
                if worksheet.max_row * worksheet.max_column > _MAX_XLSX_CELLS:
                    raise ValueError("Excel worksheet exceeds Orion's cell safety limit")
                for row_index, row in enumerate(worksheet.iter_rows(), start=1):
                    values: list[str] = []
                    for cell in row:
                        if cell.value is None:
                            continue
                        cells_seen += 1
                        if cells_seen > _MAX_XLSX_CELLS:
                            raise ValueError("Excel workbook exceeds Orion's cell safety limit")
                        values.append(f"{cell.coordinate}={cell.value}")
                    if not values:
                        continue
                    text = "\t".join(values)
                    total_characters += len(text)
                    if total_characters > _MAX_EXTRACTED_CHARACTERS:
                        raise ValueError("Excel extracted text exceeds Orion's safety limit")
                    sections.append(
                        ParsedSection(
                            text=text,
                            section=f"Sheet {worksheet.title} · row {row_index}",
                        )
                    )
        finally:
            workbook.close()
        return _parsed_document(sections)


def _validate_office_archive(content: bytes) -> None:
    """Reuse integration preflight and add traversal/compression-ratio checks for ingestion."""
    try:
        _office_zip_preflight(content)
    except DocumentCodecError as error:
        raise ValueError(error.message) from error
    try:
        with ZipFile(BytesIO(content)) as archive:
            for member in archive.infolist():
                raw_name = member.filename
                path = PurePosixPath(raw_name.replace("\\", "/"))
                if raw_name.startswith(("/", "\\")) or ".." in path.parts:
                    raise ValueError("Office document contains an unsafe archive path")
                if member.file_size and member.compress_size == 0:
                    raise ValueError("Office document has an unsafe compression ratio")
                if member.compress_size and (
                    member.file_size / member.compress_size > _MAX_ZIP_COMPRESSION_RATIO
                ):
                    raise ValueError("Office document has an unsafe compression ratio")
    except (BadZipFile, OSError) as error:
        raise ValueError("Office document is invalid or corrupt") from error


def _parsed_document(sections: list[ParsedSection]) -> ParsedDocument:
    visible = tuple(section for section in sections if section.text.strip())
    if not visible:
        raise ValueError("Document has no readable text")
    text = "\n\n".join(section.text for section in visible)
    return ParsedDocument(text=text, sections=visible)
