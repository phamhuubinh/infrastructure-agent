"""Bounded, format-aware document reads and structured edits for Linux tools."""

from __future__ import annotations

from io import BytesIO
from pathlib import PurePosixPath
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from openpyxl import load_workbook
from openpyxl.utils.cell import range_boundaries
from openpyxl.workbook.workbook import Workbook

MAX_DOCUMENT_BYTES = 4 * 1024 * 1024
MAX_DOCUMENT_ITEMS = 100
MAX_DOCUMENT_CELLS = 200
MAX_EDIT_OPERATIONS = 32
MAX_EDIT_TEXT_CHARS = 16_000


class DocumentCodecError(RuntimeError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


def document_format(path: str) -> str:
    suffix = PurePosixPath(path).suffix.lower()
    if suffix in {".doc", ".docm", ".xls", ".xlsm"}:
        raise DocumentCodecError(
            "unsupported_format", "This legacy document format is not supported."
        )
    if suffix == ".docx":
        return "docx"
    if suffix == ".xlsx":
        return "xlsx"
    return "text"


def read_document(
    path: str,
    payload: bytes,
    *,
    cursor: int = 0,
    limit: int = MAX_DOCUMENT_ITEMS,
    sheet: str | None = None,
    cell_range: str | None = None,
) -> dict[str, object]:
    fmt = document_format(path)
    _bounded_payload(payload)
    if cursor < 0 or limit < 1 or limit > MAX_DOCUMENT_ITEMS:
        raise DocumentCodecError("invalid_input", "Document paging is outside its allowed bounds.")
    if fmt == "text":
        if sheet or cell_range:
            raise DocumentCodecError(
                "invalid_input", "Text documents do not have sheets or cell ranges."
            )
        text = _decode_text(payload)
        return {
            "format": fmt,
            "cursor": cursor,
            "items": [{"index": cursor, "text": text[cursor : cursor + MAX_EDIT_TEXT_CHARS]}],
            "next_cursor": cursor + MAX_EDIT_TEXT_CHARS
            if cursor + MAX_EDIT_TEXT_CHARS < len(text)
            else None,
        }
    if fmt == "docx":
        if sheet or cell_range:
            raise DocumentCodecError(
                "invalid_input", "Word documents do not have sheets or cell ranges."
            )
        document = _load_docx(payload)
        items = _docx_items(document)
        window = items[cursor : cursor + limit]
        return {
            "format": fmt,
            "cursor": cursor,
            "items": window,
            "next_cursor": cursor + len(window) if cursor + len(window) < len(items) else None,
        }
    workbook = _load_xlsx(payload)
    if sheet is None:
        return {
            "format": fmt,
            "sheets": [
                {"name": item.title, "max_row": item.max_row, "max_column": item.max_column}
                for item in workbook.worksheets[:MAX_DOCUMENT_ITEMS]
            ],
        }
    if sheet not in workbook.sheetnames:
        raise DocumentCodecError("not_found", "Workbook sheet was not found.")
    worksheet = workbook[sheet]
    minimum_col, minimum_row, maximum_col, maximum_row = _xlsx_bounds(worksheet, cell_range)
    cells: list[dict[str, object]] = []
    for row in worksheet.iter_rows(
        min_row=minimum_row,
        max_row=maximum_row,
        min_col=minimum_col,
        max_col=maximum_col,
    ):
        for cell in row:
            cells.append({"cell": cell.coordinate, "value": _cell_value(cell.value)})
            if len(cells) == MAX_DOCUMENT_CELLS:
                return {"format": fmt, "sheet": sheet, "range": cell_range, "cells": cells}
    return {"format": fmt, "sheet": sheet, "range": cell_range, "cells": cells}


def edit_document(
    path: str, payload: bytes, operations: list[dict[str, object]]
) -> tuple[bytes, dict[str, object]]:
    fmt = document_format(path)
    _bounded_payload(payload)
    if not operations or len(operations) > MAX_EDIT_OPERATIONS:
        raise DocumentCodecError(
            "invalid_input", "Document edits require a bounded non-empty operation list."
        )
    if fmt == "text":
        result = _edit_text(_decode_text(payload), operations)
        return result.encode("utf-8"), {"format": fmt, "operations": len(operations)}
    if fmt == "docx":
        document = _load_docx(payload)
        _edit_docx(document, operations)
        output = BytesIO()
        document.save(output)
        edited = output.getvalue()
        _load_docx(edited)
        return edited, {"format": fmt, "operations": len(operations)}
    workbook = _load_xlsx(payload)
    _edit_xlsx(workbook, operations)
    output = BytesIO()
    workbook.save(output)
    edited = output.getvalue()
    _load_xlsx(edited)
    return edited, {"format": fmt, "operations": len(operations)}


def verify_document_edit(
    path: str, payload: bytes, operations: list[dict[str, object]]
) -> dict[str, object]:
    """Confirm every requested operation is observable after a write without exposing bytes."""
    fmt = document_format(path)
    _bounded_payload(payload)
    if fmt == "text":
        text = _decode_text(payload)
        for operation in operations:
            if (
                operation.get("kind") != "replace_text"
                or str(operation.get("new_text", "")) not in text
            ):
                raise DocumentCodecError("verification_failed", "Text edit could not be verified.")
    elif fmt == "docx":
        document = _load_docx(payload)
        for operation in operations:
            kind = operation.get("kind")
            if kind == "set_paragraph":
                index, text = _index(operation, "paragraph_index"), _text(operation, "text")
                if index >= len(document.paragraphs) or document.paragraphs[index].text != text:
                    raise DocumentCodecError(
                        "verification_failed", "Word paragraph edit could not be verified."
                    )
            elif kind == "set_table_cell":
                table, row, column = (
                    _index(operation, "table_index"),
                    _index(operation, "row"),
                    _index(operation, "column"),
                )
                text = _text(operation, "text")
                if (
                    table >= len(document.tables)
                    or row >= len(document.tables[table].rows)
                    or column >= len(document.tables[table].rows[row].cells)
                    or document.tables[table].cell(row, column).text != text
                ):
                    raise DocumentCodecError(
                        "verification_failed", "Word table edit could not be verified."
                    )
            else:
                raise DocumentCodecError(
                    "invalid_input", "Edit operation does not match this document format."
                )
    else:
        workbook = _load_xlsx(payload)
        for operation in operations:
            kind, sheet, cell = operation.get("kind"), operation.get("sheet"), operation.get("cell")
            if (
                kind not in {"set_cell", "set_formula"}
                or not isinstance(sheet, str)
                or not isinstance(cell, str)
            ):
                raise DocumentCodecError(
                    "invalid_input", "Edit operation does not match this document format."
                )
            _validate_cell(cell)
            if sheet not in workbook.sheetnames or workbook[sheet][cell].value != operation.get(
                "value", operation.get("formula")
            ):
                raise DocumentCodecError(
                    "verification_failed", "Workbook edit could not be verified."
                )
    return {"status": "verified", "format": fmt, "operations": len(operations)}


def _bounded_payload(payload: bytes) -> None:
    if len(payload) > MAX_DOCUMENT_BYTES:
        raise DocumentCodecError("too_large", "Document exceeds Orion's safe read limit.")


def _decode_text(payload: bytes) -> str:
    try:
        return payload.decode("utf-8")
    except UnicodeDecodeError as error:
        raise DocumentCodecError(
            "unsupported_content", "Text editing supports UTF-8 files only."
        ) from error


def _load_docx(payload: bytes) -> DocxDocument:
    try:
        return Document(BytesIO(payload))
    except Exception as error:  # python-docx normalizes several zip/XML errors.
        raise DocumentCodecError(
            "invalid_document", "Word document is invalid or corrupt."
        ) from error


def _load_xlsx(payload: bytes) -> Workbook:
    try:
        return load_workbook(BytesIO(payload), data_only=False, keep_vba=False)
    except Exception as error:  # openpyxl normalizes several zip/XML errors.
        raise DocumentCodecError(
            "invalid_document", "Excel workbook is invalid or corrupt."
        ) from error


def _docx_items(document: DocxDocument) -> list[dict[str, object]]:
    items: list[dict[str, object]] = [
        {
            "kind": "paragraph",
            "paragraph_index": index,
            "text": paragraph.text[:MAX_EDIT_TEXT_CHARS],
        }
        for index, paragraph in enumerate(document.paragraphs)
    ]
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                items.append(
                    {
                        "kind": "table_cell",
                        "table_index": table_index,
                        "row": row_index,
                        "column": column_index,
                        "text": cell.text[:MAX_EDIT_TEXT_CHARS],
                    }
                )
    return items


def _xlsx_bounds(worksheet: Any, cell_range: str | None) -> tuple[int, int, int, int]:
    if cell_range is None:
        maximum_row = min(max(worksheet.max_row, 1), MAX_DOCUMENT_CELLS)
        maximum_column = min(max(worksheet.max_column, 1), MAX_DOCUMENT_CELLS)
        return 1, 1, maximum_column, maximum_row
    try:
        minimum_col, minimum_row, maximum_col, maximum_row = range_boundaries(cell_range)
    except ValueError as error:
        raise DocumentCodecError("invalid_input", "Excel range is invalid.") from error
    if None in {minimum_col, minimum_row, maximum_col, maximum_row}:
        raise DocumentCodecError("invalid_input", "Excel range is invalid.")
    assert minimum_col is not None
    assert minimum_row is not None
    assert maximum_col is not None
    assert maximum_row is not None
    if (maximum_col - minimum_col + 1) * (maximum_row - minimum_row + 1) > MAX_DOCUMENT_CELLS:
        raise DocumentCodecError("invalid_input", "Excel range exceeds Orion's cell limit.")
    return minimum_col, minimum_row, maximum_col, maximum_row


def _edit_text(text: str, operations: list[dict[str, object]]) -> str:
    for operation in operations:
        if operation.get("kind") != "replace_text":
            raise DocumentCodecError(
                "invalid_input", "Edit operation does not match this document format."
            )
        old, new = _text(operation, "old_text"), _text(operation, "new_text")
        matches = text.count(old)
        if matches == 0:
            raise DocumentCodecError("not_found", "Text to replace was not found.")
        if matches > 1:
            raise DocumentCodecError("ambiguous", "Text to replace matches more than once.")
        text = text.replace(old, new, 1)
    return text


def _edit_docx(document: DocxDocument, operations: list[dict[str, object]]) -> None:
    for operation in operations:
        kind = operation.get("kind")
        if kind == "set_paragraph":
            index = _index(operation, "paragraph_index")
            if index >= len(document.paragraphs):
                raise DocumentCodecError("not_found", "Word paragraph was not found.")
            document.paragraphs[index].text = _text(operation, "text")
        elif kind == "set_table_cell":
            table, row, column = (
                _index(operation, "table_index"),
                _index(operation, "row"),
                _index(operation, "column"),
            )
            if (
                table >= len(document.tables)
                or row >= len(document.tables[table].rows)
                or column >= len(document.tables[table].rows[row].cells)
            ):
                raise DocumentCodecError("not_found", "Word table cell was not found.")
            document.tables[table].cell(row, column).text = _text(operation, "text")
        else:
            raise DocumentCodecError(
                "invalid_input", "Edit operation does not match this document format."
            )


def _edit_xlsx(workbook: Workbook, operations: list[dict[str, object]]) -> None:
    for operation in operations:
        kind, sheet, cell = operation.get("kind"), operation.get("sheet"), operation.get("cell")
        if (
            kind not in {"set_cell", "set_formula"}
            or not isinstance(sheet, str)
            or sheet not in workbook.sheetnames
            or not isinstance(cell, str)
        ):
            raise DocumentCodecError(
                "invalid_input", "Edit operation does not match this document format."
            )
        _validate_cell(cell)
        value = operation.get("value") if kind == "set_cell" else operation.get("formula")
        if kind == "set_formula" and (not isinstance(value, str) or not value.startswith("=")):
            raise DocumentCodecError("invalid_input", "Excel formulas must begin with '='.")
        if not isinstance(value, (str, int, float, bool)) and value is not None:
            raise DocumentCodecError("invalid_input", "Excel cell value is not supported.")
        workbook[sheet][cell].value = value


def _index(operation: dict[str, object], key: str) -> int:
    value = operation.get(key)
    if not isinstance(value, int) or isinstance(value, bool) or value < 0:
        raise DocumentCodecError("invalid_input", "Document index is invalid.")
    return value


def _text(operation: dict[str, object], key: str) -> str:
    value = operation.get(key)
    if not isinstance(value, str) or len(value) > MAX_EDIT_TEXT_CHARS:
        raise DocumentCodecError("invalid_input", "Document text is invalid or exceeds its limit.")
    return value


def _validate_cell(cell: str) -> None:
    try:
        minimum_col, minimum_row, maximum_col, maximum_row = range_boundaries(cell)
    except ValueError as error:
        raise DocumentCodecError("invalid_input", "Excel cell is invalid.") from error
    if minimum_col != maximum_col or minimum_row != maximum_row:
        raise DocumentCodecError("invalid_input", "Excel edit requires one cell.")


def _cell_value(value: object) -> str | int | float | bool | None:
    return value if isinstance(value, (str, int, float, bool)) or value is None else str(value)
